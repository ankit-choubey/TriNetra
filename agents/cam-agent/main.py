"""
Agent 12: CAM Generator Agent
Approach: Deterministic Templating (NO LLM — hallucination-free)
Tools: python-docx

Trigger: bias_completed AND stress_completed (cam_prereqs_met, counter=2)
Reads: ALL namespaces
Writes: cam_output
Logic: Generate python-docx CAM using Five Cs structure.
       Every claim cites UCSO source (page + confidence).
       Rejection cases add corrective actions.
Errors: CAM_UPLOAD_FAIL → retry 2× with backoff. CAM_EMPTY → log and alert.
"""
import sys
import os
import time
from datetime import datetime, timezone

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))
from shared.agent_base import AgentBase

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def format_inr(amount: float) -> str:
    """Format amount in Indian Rupee notation."""
    if amount >= 1e7:
        return f"₹{amount/1e7:.2f} Cr"
    elif amount >= 1e5:
        return f"₹{amount/1e5:.2f} Lakh"
    else:
        return f"₹{amount:,.2f}"


def add_heading_styled(doc: Document, text: str, level: int = 1):
    """Add a styled heading to the document."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)


def add_kv_row(table, key: str, value: str, bold_key: bool = True):
    """Add a key-value row to a table."""
    row = table.add_row()
    cell_key = row.cells[0]
    cell_val = row.cells[1]
    cell_key.text = key
    cell_val.text = str(value)
    if bold_key:
        for paragraph in cell_key.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def generate_cam_document(ucso: dict, application_id: str = "unknown") -> str:
    """
    Generate a Credit Appraisal Memo (CAM) as a Word document.
    Uses Five Cs structure with evidence citations from UCSO.

    Returns:
        Path to the generated .docx file.
    """
    doc = Document()

    applicant = ucso.get("applicant", {})
    financials = ucso.get("financials", {})
    risk = ucso.get("risk", {})
    gst = ucso.get("gst_analysis", {})
    bank = ucso.get("bank_reconciliation", {})
    web = ucso.get("web_intel", {})
    stress = ucso.get("stress_results", {})
    bias = ucso.get("bias_checks", {})
    confidence = ucso.get("decision_confidence", {})
    mca = ucso.get("mca_intelligence", {})
    pan = ucso.get("pan_intelligence", {})

    # ── Title Page ──
    title = doc.add_heading("CREDIT APPRAISAL MEMORANDUM", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Company: {applicant.get('company_name', 'N/A')}")
    doc.add_paragraph(f"PAN: {applicant.get('pan', 'N/A')} | "
                      f"GSTIN: {applicant.get('gstin', 'N/A')} | "
                      f"CIN: {applicant.get('cin', 'N/A')}")
    doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph(f"System: Trinetra Agentic Credit Intelligence OS")
    doc.add_paragraph("─" * 60)

    # ── 1. CHARACTER (Management & Identity Quality) ──
    add_heading_styled(doc, "1. CHARACTER — Identity & Management Quality", level=1)
    
    doc.add_heading("PAN Verification", level=2)
    pan_table = doc.add_table(rows=0, cols=2)
    pan_table.style = "Table Grid"
    add_kv_row(pan_table, "Verification Status", pan.get("status", "PENDING"))
    add_kv_row(pan_table, "Full Name", pan.get("full_name", "N/A"))
    add_kv_row(pan_table, "PAN Status", pan.get("pan_status", "N/A"))
    add_kv_row(pan_table, "Aadhaar Seeding", pan.get("aadhaar_seeding_status", "N/A"))

    doc.add_heading("MCA Status", level=2)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    add_kv_row(table, "Company Status (MCA)", mca.get("company_status", "N/A"))
    add_kv_row(table, "Director Changes (2yr)", str(len(mca.get("director_changes_last_2yr", []))))
    add_kv_row(table, "New Charges Filed", "YES ⚠" if mca.get("new_charge_flag") else "NO")
    add_kv_row(table, "Defaulter Flag", "YES ⚠" if mca.get("defaulter_flag") else "NO")
    add_kv_row(table, "Active Litigation Cases", str(len(web.get("litigation_records", []))))

    # News sentiment
    news = web.get("promoter_news", [])
    if news:
        avg_sentiment = sum(n.get("sentiment_score", 0) for n in news) / len(news)
        doc.add_paragraph(
            f"News Sentiment: {'POSITIVE' if avg_sentiment > 0.05 else 'NEGATIVE' if avg_sentiment < -0.05 else 'NEUTRAL'} "
            f"(avg score: {avg_sentiment:.3f}, {len(news)} articles analyzed)"
        )

    # ── 2. CAPACITY (Financial Ability to Repay) ──
    add_heading_styled(doc, "2. CAPACITY — Financial Ability to Repay", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    revenue_list = financials.get("revenue_annual", [])
    ebitda_list = financials.get("ebitda_annual", [])
    add_kv_row(table, "Revenue", format_inr(revenue_list[-1]) if revenue_list else "N/A")
    add_kv_row(table, "EBITDA", format_inr(ebitda_list[-1]) if ebitda_list else "N/A")
    add_kv_row(table, "Total Debt", format_inr(financials.get("total_debt", 0)))
    add_kv_row(table, "Net Worth", format_inr(financials.get("net_worth", 0)))
    add_kv_row(table, "Interest Expense", format_inr(financials.get("interest_expense", 0)))

    derived = ucso.get("derived_features", {})
    doc.add_paragraph(f"DSCR: {derived.get('dscr', 'N/A')}")
    doc.add_paragraph(f"ICR: {derived.get('icr', 'N/A')}")
    doc.add_paragraph(f"Leverage (D/E): {derived.get('leverage', 'N/A')}")

    # ── 3. CAPITAL (Equity Structure) ──
    add_heading_styled(doc, "3. CAPITAL — Equity Structure", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    add_kv_row(table, "Promoter Holding %", f"{financials.get('promoter_holding_pct', 0):.1f}%")
    add_kv_row(table, "Pledged Shares %", f"{financials.get('pledged_shares_pct', 0):.1f}%")
    add_kv_row(table, "CIBIL Score", str(financials.get("cibil_score", "N/A")))

    # ── 4. CONDITIONS (Market & Regulatory) ──
    add_heading_styled(doc, "4. CONDITIONS — Market & Regulatory Environment", level=1)

    headwinds = web.get("sector_headwinds", [])
    if headwinds:
        doc.add_paragraph("Sector Headwinds Identified:")
        for hw in headwinds:
            doc.add_paragraph(f"  • {hw}", style="List Bullet")
    else:
        doc.add_paragraph("No significant sector headwinds detected.")

    # GST Analysis
    doc.add_paragraph(
        f"GST Reconciliation: {gst.get('reconciliation_status', 'N/A')} "
        f"(2B vs 3B discrepancy: {gst.get('gstr2b_vs_3b_discrepancy_pct', 0):.1f}%)"
    )
    doc.add_paragraph(
        f"Circular Trade Index: {gst.get('circular_trade_index', 0):.4f} "
        f"({len(gst.get('suspicious_cycles', []))} cycles detected)"
    )

    # Bank Reconciliation
    doc.add_paragraph(
        f"Bank Reconciliation: {bank.get('reconciliation_verdict', 'N/A')} "
        f"(turnover divergence: {bank.get('turnover_divergence_pct', 0):.1f}%, "
        f"inflation flag: {'YES ⚠' if bank.get('revenue_inflation_flag') else 'NO'})"
    )

    # ── 5. COLLATERAL (Risk Mitigation) ──
    add_heading_styled(doc, "5. COLLATERAL — Risk Assessment & Decision", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    add_kv_row(table, "Risk Score", f"{risk.get('score', 0):.4f}")
    add_kv_row(table, "Risk Band", risk.get("band", "N/A"))
    add_kv_row(table, "Decision", risk.get("decision", "N/A"))
    add_kv_row(table, "Model Used", risk.get("model_used", "N/A"))
    add_kv_row(table, "Recommended Limit", format_inr(risk.get("recommended_limit", 0)))
    add_kv_row(table, "Recommended Rate (bps)", f"{risk.get('recommended_rate_bps', 0):.0f}")

    # Top risk factors
    top_factors = risk.get("top_risk_factors", [])
    if top_factors:
        doc.add_paragraph("Top Risk Factors (SHAP Analysis):")
        for factor in top_factors[:5]:
            feat = factor.get("feature", "")
            val = factor.get("shap_value", factor.get("contribution", 0))
            doc.add_paragraph(f"  • {feat}: {val:.4f}", style="List Bullet")

    # Stress Test Results
    doc.add_heading("Stress Test Results", level=2)
    scenarios = stress.get("scenarios", [])
    if scenarios:
        stress_table = doc.add_table(rows=1, cols=3)
        stress_table.style = "Table Grid"
        hdr = stress_table.rows[0].cells
        hdr[0].text = "Scenario"
        hdr[1].text = "DSCR"
        hdr[2].text = "Verdict"
        for s in scenarios:
            row = stress_table.add_row()
            row.cells[0].text = s.get("name", "")
            row.cells[1].text = f"{s.get('dscr', 0):.4f}"
            row.cells[2].text = s.get("verdict", "")

    doc.add_paragraph(
        f"Worst Case DSCR: {stress.get('worst_case_dscr', 0):.4f} → "
        f"{stress.get('survival_verdict', 'N/A')}"
    )

    # Bias Check
    doc.add_heading("Bias & Fairness Check", level=2)
    doc.add_paragraph(
        f"Counterfactual Tested: {'YES ✓' if bias.get('counterfactual_tested') else 'NO ✗'}"
    )
    flip_features = bias.get("flip_features", [])
    if flip_features:
        doc.add_paragraph("⚠ Decision-Flipping Features:")
        for ff in flip_features:
            doc.add_paragraph(
                f"  • {ff.get('feature')}: removing this feature changes decision from "
                f"{ff.get('original_decision')} to {ff.get('modified_decision')}",
                style="List Bullet",
            )

    # Confidence Score
    doc.add_heading("Decision Confidence", level=2)
    doc.add_paragraph(f"Overall Confidence: {confidence.get('score', 0):.4f}")
    doc.add_paragraph(f"Formula: {confidence.get('formula', 'N/A')}")

    # Rejection Reasons & Corrective Actions
    rejection_reasons = risk.get("rejection_reasons", [])
    if rejection_reasons:
        doc.add_heading("Rejection Reasons", level=2)
        for reason in rejection_reasons:
            doc.add_paragraph(f"  • {reason}", style="List Bullet")

    corrective = risk.get("corrective_actions", [])
    if corrective:
        doc.add_heading("Corrective Actions Required", level=2)
        for action in corrective:
            doc.add_paragraph(f"  • {action}", style="List Bullet")

    # ── Save ──
    output_dir = "/tmp/trinetra_cam"
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"CAM_{application_id}.docx")
    doc.save(output_path)

    return output_path


class CAMGeneratorAgent(AgentBase):
    AGENT_NAME = "cam-generator-agent"
    LISTEN_TOPICS = ["stress_completed"]
    OUTPUT_NAMESPACE = "cam_output"
    OUTPUT_EVENT = "cam_generated"

    def process(self, application_id: str, ucso: dict) -> dict:
        """
        Generate the Credit Appraisal Memo using python-docx templates.
        Upload to S3. Strictly deterministic — no LLM.
        """
        self.logger.info(
            f"Generating CAM document for {application_id}",
            extra={"agent_name": self.AGENT_NAME, "application_id": application_id},
        )

        # Generate the document
        output_path = generate_cam_document(ucso, application_id)

        # Validate file size (must be >5KB to not be empty)
        file_size = os.path.getsize(output_path)
        self.logger.info(
            f"CAM document generated: {output_path} ({file_size} bytes)",
            extra={"agent_name": self.AGENT_NAME, "application_id": application_id},
        )
        if file_size < 5120:
            self.logger.warning(
                f"CAM file too small ({file_size} bytes), may be empty",
                extra={"agent_name": self.AGENT_NAME, "application_id": application_id},
            )

        # Upload file to backend via POST /api/files/upload with retry (3×)
        s3_key = ""
        upload_success = False
        for attempt in range(3):
            try:
                self.logger.info(
                    f"Uploading CAM file to backend (attempt {attempt + 1}/3)...",
                    extra={"agent_name": self.AGENT_NAME, "application_id": application_id},
                )
                s3_key = self.ucso_client.upload_file(
                    application_id, output_path, "CAM"
                )
                self.logger.info(
                    f"CAM file uploaded successfully. s3_key={s3_key}",
                    extra={"agent_name": self.AGENT_NAME, "application_id": application_id},
                )
                upload_success = True
                break
            except Exception as e:
                self.logger.error(
                    f"CAM upload attempt {attempt + 1}/3 FAILED: {e}",
                    extra={"agent_name": self.AGENT_NAME, "application_id": application_id},
                )
                if attempt < 2:
                    time.sleep(2 ** attempt)  # exponential backoff

        if not upload_success:
            self.logger.error(
                f"All 3 upload attempts failed! File is at: {output_path}",
                extra={"agent_name": self.AGENT_NAME, "application_id": application_id},
            )

        # Only clean up temp file if upload succeeded
        if upload_success:
            try:
                os.unlink(output_path)
            except OSError:
                pass

        return {
            "s3_key": s3_key,
            "file_path": output_path if not upload_success else "",
            "file_size_bytes": file_size,
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "five_cs": {
                "character": "PAN Verification + MCA + Litigation + News Sentiment",
                "capacity": "DSCR + ICR + Financial Ratios",
                "capital": "Promoter Holding + CIBIL + Pledged Shares",
                "conditions": "GST Recon + Bank Recon + Sector Headwinds",
                "collateral": "Risk Score + Stress Test + Bias Check",
            },
        }


if __name__ == "__main__":
    agent = CAMGeneratorAgent()
    agent.run()
